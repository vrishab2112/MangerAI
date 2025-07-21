# app.py
import os
from pathlib import Path
import streamlit as st
from dotenv import load_dotenv
from agents.deploy_agent import parse_and_save_coder_output
import json
import re
import time
from docx import Document

# Load environment variables from .env
load_dotenv()
from agents.requirements_agent import analyze_requirements
from agents.design_agent import create_design
from agents.coder_agent import generate_code

# Initialize session states
if 'server_process' not in st.session_state:
    st.session_state.server_process = None
if 'current_report' not in st.session_state:
    st.session_state.current_report = None
if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False

def generate_word_report(agent_name, content):
    """
    Append agent output to a shared Word document report.
    This function takes output from each agent (Requirements, Design, and Coder agents)
    and adds it to a single Word document, organizing the content by agent.
    
    Parameters:
        agent_name (str): Name of the agent (e.g. "Requirements Agent", "Design Agent")
        content (str): The output/report content from the agent
        
    Returns:
        str: Path to the generated Word document
    """
    report_path = "agent_reports.docx"
    
    # Create new document or load existing one if it exists
    if os.path.exists(report_path):
        doc = Document(report_path)
    else:
        doc = Document()
        doc.add_heading('Multi-Agent System Report', 0)
        doc.add_paragraph('This document contains the combined outputs from all agents in the system.')
    
    # Add a section for this agent's report
    doc.add_heading(f'{agent_name} Report', level=1)
    doc.add_paragraph(content)
    doc.add_paragraph('\n') # Add spacing between sections
    
    # Save the updated document
    doc.save(report_path)
    st.session_state.current_report = report_path
    st.session_state.report_generated = True
    
    return report_path

def download_report():
    """
    Creates a download button in the Streamlit interface for the Word report.
    When clicked, allows downloading the combined report containing all agent outputs.
    """
    if st.session_state.report_generated and os.path.exists(st.session_state.current_report):
        with open(st.session_state.current_report, 'rb') as f:
            report_data = f.read()
        return report_data
    return None

# Streamlit UI
st.set_page_config(page_title="Multi-Agent Code Generator")
st.title("Multi-Agent Code Generation and Deployment System")
st.write(
    "Enter a high-level prompt (e.g. *'Build me a website'* or *'Build me a snake game'*). "
    "Select models for each agent and click **Generate** to run the pipeline."
)

# Input prompt
prompt = st.text_input("Enter your prompt:")

# Model selection options
model_options = [
    "llama3-8b-8192",
    "Groq/Llama-3-Groq-8B-Tool-Use",
    "Groq/Llama-3-Groq-70B-Tool-Use",
    "gpt-3.5-turbo"
]

col1, col2, col3 = st.columns(3)
with col1:
    req_model = st.selectbox("Requirements Agent Model:", model_options, index=0)
with col2:
    design_model = st.selectbox("Design Agent Model:", model_options, index=0)
with col3:
    coder_model = st.selectbox("Coder Agent Model:", model_options, index=0)

# Button to trigger the pipeline
if st.button("Generate"):
    if not prompt:
        st.error("Please enter a prompt.")
    else:
        # Reset report state for new generation
        st.session_state.report_generated = False
        
        # 1. Requirements Analysis
        st.subheader("1. Requirements Analysis")
        try:
            requirements = analyze_requirements(prompt, req_model)
            st.text(requirements)
            generate_word_report("Requirements Agent", requirements)
        except Exception as e:
            st.error(f"Error in Requirements Agent: {e}")
            requirements = ""

        # 2. Design Specification
        if requirements:
            st.subheader("2. Design Specification")
            try:
                design = create_design(requirements, design_model)
                st.text(design)
                generate_word_report("Design Agent", design)
            except Exception as e:
                st.error(f"Error in Design Agent: {e}")
                design = ""
        else:
            design = ""

        # 3. Code Generation
        if design:
            st.subheader("3. Code Generation")
            try:
                # Stop any existing server process
                if st.session_state.server_process:
                    st.session_state.server_process.terminate()
                    st.session_state.server_process = None
                
                code_output = generate_code(design, coder_model)
                st.markdown(code_output)
                generate_word_report("Code Generation Agent", code_output)
                
                # 4. Save and deploy the generated code
                st.subheader("4. Deployment")
                url, proc = parse_and_save_coder_output(code_output)
                
                # Show success message and file location
                deployment_status = f"Code files have been saved to the 'deployed_app' directory\n"
                if url and proc:
                    st.session_state.server_process = proc
                    deployment_status += f"Application launched at: {url}"
                    st.success(f"🚀 Application launched! Access it at: {url}")
                    
                    # Add a button to stop the server
                    if st.button("Stop Server"):
                        if st.session_state.server_process:
                            st.session_state.server_process.terminate()
                            st.session_state.server_process = None
                            st.info("Server stopped")
                            deployment_status += "\nServer stopped by user"
                
                # Generate report for Deployment Agent
                generate_word_report("Deployment Agent", deployment_status)
                
                # Create download buttons for code and report
                col1, col2 = st.columns(2)
                with col1:
                    if os.path.exists("deployed_app"):
                        import shutil
                        zip_path = shutil.make_archive("deployed_app", 'zip', "deployed_app")
                        with open(zip_path, 'rb') as f:
                            st.download_button(
                                label="Download Code as ZIP",
                                data=f,
                                file_name="deployed_app.zip",
                                mime="application/zip"
                            )
                
            except Exception as e:
                st.error(f"Error in Code Generation/Deployment: {e}")
                code_output = ""
        else:
            code_output = ""

# Add the report download button outside the generate button block
if st.session_state.report_generated:
    report_data = download_report()
    if report_data:
        st.download_button(
            label="Download Combined Agent Reports",
            data=report_data,
            file_name="agent_reports.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Cleanup server process when the app is closed
def cleanup():
    if st.session_state.server_process:
        st.session_state.server_process.terminate()
        st.session_state.server_process = None

# Register the cleanup function
import atexit
atexit.register(cleanup)

def parse_and_save_coder_output(code_output):
    """
    Parse the code output and save it to the workspace using the deployment agent.
    """
    try:
        from deployment_agent import parse_code_blocks, save_to_workspace
        
        # Parse the code blocks
        code_blocks = parse_code_blocks(code_output)
        if not code_blocks:
            st.error("No code blocks with filenames found in the output.")
            return
        
        # Save to workspace
        workspace_dir = save_to_workspace(code_blocks)
        st.success(f"Code files have been saved to the '{workspace_dir}' directory")
        
        # Create a download button for the directory
        if st.button("Download Code"):
            import shutil
            zip_path = shutil.make_archive(workspace_dir, 'zip', workspace_dir)
            with open(zip_path, 'rb') as f:
                st.download_button(
                    label="Download ZIP",
                    data=f,
                    file_name=f"{workspace_dir}.zip",
                    mime="application/zip"
                )
    except Exception as e:
        st.error(f"Error saving code: {e}")

